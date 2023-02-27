import streamlit as st
import random
import string
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from sqlalchemy.orm import sessionmaker
from sqlalchemy.ext.declarative import declarative_base
import stripe
import datetime
import pandas as pd
from sqlalchemy import create_engine, Column, Integer, String, Float, DateTime,Boolean
import sqlalchemy
import re
import time
import hashlib
from datetime import datetime
from PIL import Image

import matplotlib.pyplot as plt
import streamlit as st
from wordcloud import WordCloud,ImageColorGenerator
from pytrends.request import TrendReq

# import nltk
# from nltk.sentiment.vader import SentimentIntensityAnalyzer
# nltk.downloader.download('vader_lexicon')
from collections import Counter
from datetime import date
import openai
import pandas as pd
import requests
import os
import streamlit as st
from PIL import Image
from io import BytesIO
import docx


st.set_page_config(page_title="OpenAI Web App", page_icon="🧊", layout="wide")


# Read the data from the Excel file
df = pd.read_excel('prompt_list.xlsx')
# Set page configuration

#Hide streamlit logo and hamberger!
hide_streamlit_style = """
            <style>
            #MainMenu {visibility: hidden;}
            footer {visibility: hidden;}
            </style>
            """
st.markdown(hide_streamlit_style, unsafe_allow_html=True)

## Get API key

openai.organization = "org-ixe83K3WLxi5O4MhZwLa0ABR"
openai.api_key = "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"
trendslist = []
google_trend_list=[]

prompt_list=[]

def openai_general(prompt):

    # Set the OpenAI model and generate text based on the prompt
    model = "text-davinci-003"
    completions = openai.Completion.create(engine=model, prompt=prompt, max_tokens=2048, n=1, stop=None, temperature=0.5)

    # Extract the generated text from the API response
    generated_text_openai = completions.choices[0].text

    return generated_text_openai


from urllib.parse import quote


def DalleImage(input,content):

    now = datetime.now()
    timestamp=str(now.year)+str(now.month)+str(now.day)+str(now.hour)+str(now.minute)+str(now.second)

    # prompt=message
    prompt = content
    style=" "
    effect=" "
    final_prompt=style+effect+prompt
    final_prompt_adj=final_prompt
    print("final prompt adj",final_prompt_adj)
    image_resp = openai.Image.create(prompt=final_prompt_adj, n=1, size='512x512')
    image_url = image_resp['data'][0]['url']

    # print(image_url)

    img_data = requests.get(image_url).content
    with open(timestamp+'.jpg', 'wb') as handler:
        handler.write(img_data)

    # Create a new docx document and add the text and image to it
    doc = docx.Document()
    doc.add_paragraph(input)
    doc.add_paragraph(content)
    doc.add_picture(timestamp + '.jpg')

    # Read the image from disk and display it
    img = Image.open(timestamp+'.jpg')
    st.image(img, caption='Generated Image', use_column_width=True)
    # Save the document to a local folder
    sentence = prompt.translate(str.maketrans("", "", string.punctuation))
    # remove non-Chinese characters
    sentence = re.sub(r'[^\u4e00-\u9fff]+', '', sentence)

    doc.save(sentence[:15]+'.docx')


def googletrend():
    timeframe = 'today 1-m'  # past 1 month
    global aoa, google_trend_list, trendslist

    try:
        pytrends = TrendReq(hl='ja-JP', tz=360)
    except:
        st.error("Connection Error: Failed to connect to Google Trends. Please try again later.")
        # You can add more instructions or information for the user here if needed.

    pnlist = ['united_states', 'japan', 'hong_kong', 'united_kingdom', 'taiwan', 'singapore', 'australia']
    trendslist = []

    for l in pnlist:
        try:
            trends = pytrends.trending_searches(pn=l)
            trends_values = trends.values.tolist()
            trendslist.extend(trends_values)
        except:
            st.write("Network error, please try again later")





    for value in trendslist:
        item = str(value[0])
        item = item.replace('  ', '')
        google_trend_list.append(item)

    sorted_trends = sorted(Counter(google_trend_list).items(), key=lambda pair: pair[1], reverse=True)
    wordcloud()
    st.image(exportName + ".png", caption=exportName)
    st.write('### Sorted Trends')

    # Create a paragraph with clickable hyperlinks
    trend_paragraphs = []
    for trend, count in sorted_trends:
        link = f'https://www.google.com/search?q={trend}'
        trend_paragraphs.append(f'<a href="{link}" target="_blank">{trend}</a>')

    st.write('<br>'.join(trend_paragraphs), unsafe_allow_html=True)


    aoa = trendslist


    return trendslist, google_trend_list




def color_func(word, font_size, position, orientation, random_state=None, **kwargs):
    colors = ["#69ACDE","#87AFC5","#ffb8e8","#6eafff","#FFFAFA", "#ffaea3","#e9abff","#EAE8F2", "#E2D0C3","#FCEABB","#F5FFFA","#F0F0F0","#CAE3BF","#D8BFD8"]  # Define a list of colors
    return random.choice(colors)  # Choose a random color from the list

def wordcloud():
    global google_trend_list
    now = datetime.now()
    hour = now.strftime("%H")
    minute = now.strftime("%M")
    print(google_trend_list)

    global exportName
    font = r'msyh.ttc'  # 設定字型

    # for i in imageList:

        # mask = np.array(Image.open("cloud.png"))  # 設定文字雲形狀
    unioncloud = WordCloud(background_color="black",color_func=color_func ,font_path=font,width=1600, height=1600)  # 背景顏色預設黑色,改為白色
    unioncloud.generate_from_frequencies(frequencies=Counter(google_trend_list))  # 產生文字雲

    #colormap='tab20'

    # imname=i
    # inputimage=plt.imread(imname+'.jpg')
    # img_colors=ImageColorGenerator(inputimage)

    # 產生圖片
    plt.figure(figsize=(10, 10),dpi=100)

    # unioncloud.recolor(color_func=img_colors)

    plt.imshow(unioncloud)
    plt.axis("off")
    today = date.today()
    # plt.show()



    exportName=str(today)+"-"+hour+"-"+minute
    unioncloud.to_file(exportName+".png")  # 存檔


def main():

    # Hide Streamlit logo and hamberger!
    hide_streamlit_style = """
                <style>
                #MainMenu {visibility: hidden;}
                footer {visibility: hidden;}
                </style>
                """
    st.markdown(hide_streamlit_style, unsafe_allow_html=True)

    # Define ad hoc pages
    ad_hoc_pages = {"Trend/热点": {"intro": "Global Latest Focus Keywords/全球最新热点关键词"},
                    "Ad Hoc Test": {"intro": "This is ad hoc"}}

    # Read the data from the Excel file
    df = pd.read_excel('prompt_list.xlsx')

    # Create a dropdown menu to select the category
    categories = ["All/全部", "Other/其他"] + df['Category'].unique().tolist()
    selected_category = st.sidebar.selectbox("Select a category/请选择一个类别", categories)

    # Filter the data based on the selected category
    if selected_category == "All/全部":
        filtered_df = df
    elif selected_category == "Other/其他":
        filtered_df = df[df['Category'] == "Other/其他"].append(pd.DataFrame.from_dict(ad_hoc_pages).T, ignore_index=False)
    else:
        filtered_df = df[df['Category'] == selected_category]

    # Create a dropdown menu to select the page
    if selected_category == "Other/其他":
        pages = ["---Please select/请选择---"] + list(ad_hoc_pages.keys()) + filtered_df[filtered_df['Category'] == "Other/其他"]['page'].unique().tolist()
    elif selected_category == "All/全部":
        pages = ["---Please select/请选择---"] + list(ad_hoc_pages.keys())+filtered_df['page'].unique().tolist()

    else:
        pages = ["---Please select/请选择---"] + filtered_df['page'].unique().tolist()

    selected_page = st.sidebar.selectbox("Select a page/请选择", pages)
    # selected_page = st.sidebar.selectbox("Select a page/请选择", pages, **{"width": "300px"})
    if selected_page in ad_hoc_pages:
        # Set the page title and intro for ad hoc pages
        st.title(ad_hoc_pages[selected_page]["intro"])

        # Add a button to run Google Trend
        if selected_page=="Trend/热点":
            if st.button("Generate/运行"):
                with st.spinner('Program is running, please wait.../程序正在运行中，请耐心等待...'):
                    googletrend()

    else:
        # Get the row corresponding to the selected page
        page_data = filtered_df[filtered_df['page'] == selected_page]
        if len(page_data) == 0:
            st.write(f"{selected_page}")
            return
        page_data = page_data.iloc[0]

        # Set the page title and intro
        st.title(page_data['intro'])

        # Add an input box with the specified name
        input_name = page_data['input']
        user_input = st.text_area(input_name, placeholder="Keywords can be separated by commas/关键词可以用逗号分开",height=10)

        # Add a dropdown menu to select reply type and language
        # reply_type = st.selectbox("Select reply type", ["Detailed Explanation/详细解释","Key Takeaways/关键要点"])
        language = st.selectbox("Select language", ["Chinese/中文","English/英文",  "Japanese/日语"])

        restriction = "please answers in following requirement: " + "  用以下语言" + language

        # Add a generate button
        if st.button("Generate/运行"):
            with st.spinner('Program is running, please wait.../程序正在运行中，请耐心等待...'):

                if len(page_data) > 0:
                    prompt = page_data.loc["prompt"]+". Input info is: {A}.{B}. 回答字数尽量多，内容尽量详细，但控制在输出最大字数限制范围的90%左右，避免产生不完整答复。"
                    prompt = prompt.replace("{A}", user_input).replace("{B}", restriction)

                    # Add some formatted text to the page
                    st.markdown("<span style='color: blue;'>Here is the answer to your question/生成答案:</span>",
                                unsafe_allow_html=True)

                    # Generate and display the answer
                    generated_text_openai = openai_general(prompt)
                    st.write(generated_text_openai)
                    if selected_category == "Arts & Culture/艺术和文化":
                        DalleImage(user_input, generated_text_openai)



                else:
                    print(f"No prompt found for page")


if __name__ == "__main__":
    main()

